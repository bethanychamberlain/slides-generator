"""Word document generation for study guides."""

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def format_question_for_docx(doc, q, question_num, show_answers=False):
    """Format a structured question for the Word document - clear visual hierarchy."""
    qtype = q.get("type", "open_ended")
    indent = Inches(0.3)

    if qtype == "open_ended":
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(q.get('question', ''))
        p.paragraph_format.space_after = Pt(2)
        if show_answers:
            example = q.get('example_answer', '')
            answer_p = doc.add_paragraph()
            answer_p.paragraph_format.left_indent = indent
            if example:
                run = answer_p.add_run("Example answer: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 100, 0)
                run2 = answer_p.add_run(example)
                run2.font.color.rgb = RGBColor(0, 100, 0)
            else:
                run = answer_p.add_run("[Open-ended response]")
                run.italic = True
                run.font.color.rgb = RGBColor(0, 100, 0)
            answer_p.paragraph_format.space_after = Pt(10)
        else:
            for line_idx in range(2):
                line = doc.add_paragraph()
                run = line.add_run("_" * 70)
                run.font.size = Pt(14)
                line.paragraph_format.left_indent = indent
                line.paragraph_format.space_before = Pt(0)
                line.paragraph_format.space_after = Pt(0) if line_idx < 1 else Pt(10)

    elif qtype == "short_answer":
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(q.get('prompt', ''))
        p.paragraph_format.space_after = Pt(2)
        if show_answers:
            ans_p = doc.add_paragraph()
            ans_p.paragraph_format.left_indent = indent
            run = ans_p.add_run(f"Answer: {q.get('answer', '')}")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0)
            ans_p.paragraph_format.space_after = Pt(10)
        else:
            line = doc.add_paragraph()
            run = line.add_run("_" * 70)
            run.font.size = Pt(14)
            line.paragraph_format.left_indent = indent
            line.paragraph_format.space_after = Pt(10)

    elif qtype == "fill_in_blank":
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        answer = q.get('answer', '')
        if show_answers:
            sentence = q.get('sentence', '').replace('_____', f"[{answer}]")
            p.add_run(sentence)
            p.paragraph_format.space_after = Pt(2)
            ans_p = doc.add_paragraph()
            ans_p.paragraph_format.left_indent = indent
            run = ans_p.add_run(f"Answer: {answer}")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0)
            ans_p.paragraph_format.space_after = Pt(10)
        else:
            sentence = q.get('sentence', '').replace('_____', '_' * 20)
            p.add_run(sentence)
            p.paragraph_format.space_after = Pt(10)

    elif qtype == "true_false":
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(q.get('statement', ''))
        p.paragraph_format.space_after = Pt(2)
        answer = q.get('answer', True)
        if show_answers:
            opts = doc.add_paragraph()
            opts.paragraph_format.left_indent = indent
            if answer:
                run_t = opts.add_run("● True")
                run_t.bold = True
                run_t.font.color.rgb = RGBColor(0, 100, 0)
                opts.add_run("                    ○ False")
            else:
                opts.add_run("○ True                    ")
                run_f = opts.add_run("● False")
                run_f.bold = True
                run_f.font.color.rgb = RGBColor(0, 100, 0)
            opts.paragraph_format.space_after = Pt(10)
        else:
            opts = doc.add_paragraph("○  True                    ○  False")
            opts.paragraph_format.left_indent = indent
            opts.paragraph_format.space_after = Pt(10)

    elif qtype == "multiple_choice":
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(q.get('question', ''))
        p.paragraph_format.space_after = Pt(2)

        options = q.get("options", [])
        correct = q.get("answer", "")

        if options:
            table = doc.add_table(rows=1, cols=len(options))
            table.autofit = True
            row = table.rows[0]
            for opt_idx, opt in enumerate(options):
                cell = row.cells[opt_idx]
                para = cell.paragraphs[0]
                is_correct = opt.startswith(correct) if correct else False
                if show_answers and is_correct:
                    run = para.add_run(opt)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 100, 0)
                else:
                    para.add_run(opt)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblInd = OxmlElement('w:tblInd')
            tblInd.set(qn('w:w'), str(int(indent.twips)))
            tblInd.set(qn('w:type'), 'dxa')
            tblPr.append(tblInd)
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)

            if show_answers:
                ans_p = doc.add_paragraph()
                ans_p.paragraph_format.left_indent = indent
                run = ans_p.add_run(f"Answer: {correct}")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 100, 0)
                ans_p.paragraph_format.space_after = Pt(10)
            else:
                p_after = doc.add_paragraph()
                p_after.paragraph_format.space_before = Pt(0)
                p_after.paragraph_format.space_after = Pt(10)

    elif qtype == "put_in_order":
        instruction = q.get("instruction", "Arrange in order:")
        items = q.get("items", [])
        items_text = ", ".join(items)
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(f"{instruction} {items_text}")
        p.paragraph_format.space_after = Pt(2)

        if show_answers:
            correct_order = q.get("correct_order", [])
            if correct_order and items:
                ordered = [items[oi] for oi in correct_order if oi < len(items)]
                correct_text = " → ".join(ordered)
            else:
                correct_text = ", ".join(items)
            ans_p = doc.add_paragraph()
            ans_p.paragraph_format.left_indent = indent
            run = ans_p.add_run(f"Correct order: {correct_text}")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0)
            ans_p.paragraph_format.space_after = Pt(10)
        else:
            answer = doc.add_paragraph("Order: " + "_" * 45)
            answer.paragraph_format.left_indent = indent
            answer.paragraph_format.space_after = Pt(10)

    else:
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(str(q))
        p.paragraph_format.space_after = Pt(10)


def create_docx(intro_summary, questions_by_slide, outro_summary, show_answers=False):
    """Generate a compact Word document - Times New Roman, narrow margins, no title."""
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    if intro_summary:
        p = doc.add_paragraph()
        run = p.add_run("Overview: ")
        run.bold = True
        p.add_run(intro_summary)
        p.paragraph_format.space_after = Pt(8)

    question_num = 1
    for slide_num, questions in questions_by_slide.items():
        for q in questions:
            format_question_for_docx(doc, q, question_num, show_answers)
            question_num += 1

    if outro_summary:
        p = doc.add_paragraph()
        run = p.add_run("Key Takeaways: ")
        run.bold = True
        p.add_run(outro_summary)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
