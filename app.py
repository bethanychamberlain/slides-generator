import streamlit as st
from pdf2image import convert_from_bytes
from anthropic import Anthropic
import base64
import json
import hashlib
import shutil
import csv
from pathlib import Path
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv(override=True)

# Initialize Anthropic client (uses ANTHROPIC_API_KEY env variable)
client = Anthropic()

# Model configuration
MODEL_FAST = "claude-sonnet-4-20250514"  # Cheaper, for question generation
MODEL_ADVANCED = "claude-opus-4-20250514"  # Advanced, for review and selection

# Working directory for full-size images
WORKING_IMAGES_DIR = Path(__file__).parent / "working_images"

# Cache directory for storing analysis results
CACHE_DIR = Path(__file__).parent / "analysis_cache"


def get_file_hash(file_bytes):
    """Generate a hash of file contents for cache lookup."""
    return hashlib.sha256(file_bytes).hexdigest()[:16]


def get_image_hash(image):
    """Generate a hash of an image for per-slide cache lookup."""
    buffer = BytesIO()
    # Convert to consistent format for hashing
    image.save(buffer, format="PNG")
    return hashlib.sha256(buffer.getvalue()).hexdigest()[:16]


def get_cache_path(file_hash):
    """Get the path to a cached analysis file."""
    CACHE_DIR.mkdir(exist_ok=True)
    return CACHE_DIR / f"{file_hash}.json"


def get_slide_cache_path():
    """Get the path to the per-slide image cache."""
    CACHE_DIR.mkdir(exist_ok=True)
    return CACHE_DIR / "_slide_image_cache.json"


def load_slide_cache():
    """Load the per-slide image hash cache."""
    cache_path = get_slide_cache_path()
    if cache_path.exists():
        try:
            with open(cache_path, 'r') as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return {}
    return {}


def save_slide_cache(cache):
    """Save the per-slide image hash cache."""
    cache_path = get_slide_cache_path()
    with open(cache_path, 'w') as f:
        json.dump(cache, f, indent=2)


def get_cached_questions_for_image(image):
    """Check if we have cached questions for this exact image."""
    image_hash = get_image_hash(image)
    slide_cache = load_slide_cache()
    return slide_cache.get(image_hash), image_hash


def save_questions_for_image(image_hash, questions):
    """Save questions for a specific image hash."""
    slide_cache = load_slide_cache()
    slide_cache[image_hash] = questions
    save_slide_cache(slide_cache)


def load_from_cache(file_hash):
    """Load cached analysis results if available."""
    cache_path = get_cache_path(file_hash)
    if cache_path.exists():
        try:
            with open(cache_path, 'r') as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return None
    return None


def save_to_cache(file_hash, data):
    """Save analysis results to cache."""
    cache_path = get_cache_path(file_hash)
    CACHE_DIR.mkdir(exist_ok=True)
    with open(cache_path, 'w') as f:
        json.dump(data, f, indent=2)


def clear_cache():
    """Clear all cached analysis results."""
    if CACHE_DIR.exists():
        shutil.rmtree(CACHE_DIR)
    CACHE_DIR.mkdir(exist_ok=True)


# Questions history CSV for future retrieval
QUESTIONS_CSV_PATH = Path(__file__).parent / "questions_history.csv"

CSV_COLUMNS = [
    "timestamp", "source_file", "slide_num", "question_type", "question_text",
    "answer", "example_answer", "options", "items", "correct_order", "selected"
]


def save_questions_to_csv(questions_by_slide, source_filename, selected_set=None):
    """Save all generated questions to CSV for future retrieval.

    Args:
        questions_by_slide: Dict of {slide_num: [questions]}
        source_filename: Name of the source PDF file
        selected_set: Optional set of (slide_num, index) tuples that are selected
    """
    file_exists = QUESTIONS_CSV_PATH.exists()
    timestamp = datetime.now().isoformat()

    with open(QUESTIONS_CSV_PATH, 'a', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=CSV_COLUMNS)
        if not file_exists:
            writer.writeheader()

        for slide_num, questions in questions_by_slide.items():
            for i, q in enumerate(questions):
                qtype = q.get("type", "unknown")

                # Extract the main question text based on type
                if qtype == "open_ended":
                    text = q.get("question", "")
                elif qtype == "short_answer":
                    text = q.get("prompt", "")
                elif qtype == "fill_in_blank":
                    text = q.get("sentence", "")
                elif qtype == "true_false":
                    text = q.get("statement", "")
                elif qtype == "multiple_choice":
                    text = q.get("question", "")
                elif qtype == "put_in_order":
                    text = q.get("instruction", "")
                else:
                    text = str(q)

                # Get answer based on type
                answer = ""
                if qtype in ["short_answer", "fill_in_blank"]:
                    answer = q.get("answer", "")
                elif qtype == "true_false":
                    answer = str(q.get("answer", ""))
                elif qtype == "multiple_choice":
                    answer = q.get("answer", "")

                # Check if selected
                is_selected = selected_set is None or (slide_num, i) in selected_set

                row = {
                    "timestamp": timestamp,
                    "source_file": source_filename,
                    "slide_num": slide_num,
                    "question_type": qtype,
                    "question_text": text,
                    "answer": answer,
                    "example_answer": q.get("example_answer", ""),
                    "options": json.dumps(q.get("options", [])) if q.get("options") else "",
                    "items": json.dumps(q.get("items", [])) if q.get("items") else "",
                    "correct_order": json.dumps(q.get("correct_order", [])) if q.get("correct_order") else "",
                    "selected": "1" if is_selected else "0"
                }
                writer.writerow(row)


def load_questions_from_csv(source_filename=None):
    """Load questions from CSV, optionally filtered by source file.

    Args:
        source_filename: Optional filter by source file name

    Returns:
        List of question dicts
    """
    if not QUESTIONS_CSV_PATH.exists():
        return []

    questions = []
    with open(QUESTIONS_CSV_PATH, 'r', newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            if source_filename and row["source_file"] != source_filename:
                continue
            questions.append(row)

    return questions


def ensure_working_dir():
    """Create working images directory if it doesn't exist."""
    WORKING_IMAGES_DIR.mkdir(exist_ok=True)


def cleanup_working_dir():
    """Remove all images from working directory."""
    if WORKING_IMAGES_DIR.exists():
        shutil.rmtree(WORKING_IMAGES_DIR)
    WORKING_IMAGES_DIR.mkdir(exist_ok=True)


def save_slide_image(image, slide_num):
    """Save a slide image to the working directory."""
    ensure_working_dir()
    filename = WORKING_IMAGES_DIR / f"slide_{slide_num:03d}.jpg"
    image.save(filename, format="JPEG", quality=95)
    return filename


def get_slide_image_path(slide_num):
    """Get the path to a saved slide image."""
    return WORKING_IMAGES_DIR / f"slide_{slide_num:03d}.jpg"

def image_to_base64(image):
    """Convert PIL image to base64 string."""
    buffer = BytesIO()
    image.save(buffer, format="JPEG")
    return base64.standard_b64encode(buffer.getvalue()).decode("utf-8")

QUESTION_TYPES = ["open_ended", "short_answer", "fill_in_blank", "true_false", "multiple_choice", "put_in_order"]

QUESTION_TYPE_LABELS = {
    "open_ended": "Open-ended",
    "short_answer": "Short Answer",
    "fill_in_blank": "Fill in blank",
    "true_false": "True/False",
    "multiple_choice": "Multiple Choice",
    "put_in_order": "Put in order"
}


def analyze_slide(image, slide_num, total_slides, context="content"):
    """Send slide to Claude for analysis."""
    base64_image = image_to_base64(image)

    if context == "intro":
        prompt = """Analyze this introductory slide and write a 1-2 sentence overview.
Include the topic/title and main learning objectives.
Write in plain text only - no markdown, no bullets, no special formatting.
Be concise and direct."""

        response = client.messages.create(
            model=MODEL_FAST,
            max_tokens=1000,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": base64_image}},
                    {"type": "text", "text": prompt}
                ]
            }]
        )
        return response.content[0].text

    elif context == "outro":
        prompt = """Analyze this concluding slide and write a 1-2 sentence summary of the key takeaways.
Write in plain text only - no markdown, no bullets, no special formatting.
Be concise and direct."""

        response = client.messages.create(
            model=MODEL_FAST,
            max_tokens=1000,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": base64_image}},
                    {"type": "text", "text": prompt}
                ]
            }]
        )
        return response.content[0].text

    else:
        # Content slide - generate structured questions
        prompt = """Analyze this lecture slide carefully, including any charts, graphs, or diagrams.

Generate exactly 3 questions for a student note-taking guide:
1. One OPEN-ENDED question (REQUIRED) - encouraging critical thinking about the content
2. Two additional questions - choose from the types below

PRIORITY ORDER (use higher priority types when appropriate for the content):
1. OPEN_ENDED (already required) - Critical thinking, synthesis, analysis questions
2. SHORT_ANSWER - For acronyms, definitions, key terms (e.g., "What does PAMP stand for?")
3. FILL_IN_BLANK - Key terms in context (e.g., "The _____ is the powerhouse of the cell")
4. TRUE_FALSE - Correct common misconceptions or verify understanding of facts
5. MULTIPLE_CHOICE - ONLY when there are meaningfully different conceptual alternatives (NOT for acronyms or simple definitions)
6. PUT_IN_ORDER - Processes, sequences, or steps

IMPORTANT RULES:
- NEVER use multiple choice for acronyms - use short_answer instead
- NEVER use multiple choice for simple definitions - use short_answer or fill_in_blank
- Prefer short_answer and fill_in_blank over multiple_choice when possible

Return ONLY valid JSON in this exact format (no other text):
{
  "questions": [
    {
      "type": "open_ended",
      "question": "Your open-ended question here?",
      "example_answer": "A model answer showing what a good response would include (2-3 sentences)",
      "notes_prompt": "[Your notes:]"
    },
    {
      "type": "short_answer",
      "prompt": "What does PAMP stand for?",
      "answer": "Pathogen-Associated Molecular Patterns"
    },
    {
      "type": "fill_in_blank",
      "sentence": "The _____ is responsible for energy production.",
      "answer": "mitochondria"
    }
  ]
}

Other valid question formats:
- true_false: {"type": "true_false", "statement": "...", "answer": true}
- multiple_choice: {"type": "multiple_choice", "question": "...", "options": ["A) ...", "B) ...", "C) ...", "D) ..."], "answer": "B"}
- put_in_order: {"type": "put_in_order", "instruction": "Arrange these steps in order:", "items": ["Step A", "Step B", "Step C"], "correct_order": [2, 0, 1]}

IMPORTANT: For open_ended questions, always include an "example_answer" field with a model answer based on the slide content."""

        response = client.messages.create(
            model=MODEL_FAST,
            max_tokens=1500,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": base64_image}},
                    {"type": "text", "text": prompt}
                ]
            }]
        )

        # Parse JSON response
        response_text = response.content[0].text.strip()
        # Handle potential markdown code blocks
        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]
            response_text = response_text.strip()

        try:
            return json.loads(response_text)
        except json.JSONDecodeError:
            # Fallback: return as single open-ended question
            return {"questions": [{"type": "open_ended", "question": response_text, "notes_prompt": "[Your notes:]"}]}


def verify_answers(image, questions):
    """Verify and correct answers for questions using Claude."""
    base64_image = image_to_base64(image)

    # Build verification prompt
    questions_json = json.dumps(questions, indent=2)

    prompt = f"""Look at this slide and verify the correctness of the answers for these questions.
For each question, check if the answer is correct based on the slide content.
If any answer is wrong, provide the correct answer.

Current questions and answers:
{questions_json}

Return ONLY valid JSON with the corrected questions in the same format.
If all answers are correct, return the questions unchanged.
Make sure every answer is accurate based on the slide content.
IMPORTANT: Preserve all fields including example_answer for open-ended questions. If an example_answer is missing or could be improved, add or improve it."""

    response = client.messages.create(
        model=MODEL_FAST,
        max_tokens=2000,
        messages=[{
            "role": "user",
            "content": [
                {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": base64_image}},
                {"type": "text", "text": prompt}
            ]
        }]
    )

    response_text = response.content[0].text.strip()
    if response_text.startswith("```"):
        response_text = response_text.split("```")[1]
        if response_text.startswith("json"):
            response_text = response_text[4:]
        response_text = response_text.strip()

    try:
        result = json.loads(response_text)
        # Handle both {"questions": [...]} and [...] formats
        if isinstance(result, list):
            return result
        return result.get("questions", questions)
    except json.JSONDecodeError:
        return questions  # Return original if parsing fails


def generate_example_answers(questions_by_slide, slide_images):
    """Generate example answers for open-ended questions that don't have them."""
    updated_questions = {}
    
    for slide_num, questions in questions_by_slide.items():
        updated_qs = []
        needs_generation = []
        
        for i, q in enumerate(questions):
            if q.get("type") == "open_ended" and not q.get("example_answer"):
                needs_generation.append((i, q))
            updated_qs.append(q.copy())
        
        # Generate example answers for questions that need them
        if needs_generation and slide_num <= len(slide_images):
            image = slide_images[slide_num - 1]
            base64_image = image_to_base64(image)
            
            for idx, q in needs_generation:
                prompt = f"""Look at this slide and provide a model answer (2-3 sentences) for this open-ended question:

Question: {q.get('question', '')}

Write a clear, concise example answer that a student might give based on the slide content.
Return ONLY the answer text, no additional formatting or explanation."""
                
                try:
                    response = client.messages.create(
                        model=MODEL_FAST,
                        max_tokens=500,
                        messages=[{
                            "role": "user",
                            "content": [
                                {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": base64_image}},
                                {"type": "text", "text": prompt}
                            ]
                        }]
                    )
                    updated_qs[idx]["example_answer"] = response.content[0].text.strip()
                except Exception:
                    updated_qs[idx]["example_answer"] = "[Example answer not available]"
        
        updated_questions[slide_num] = updated_qs
    
    return updated_questions


def review_and_select_questions(all_questions, total_slides):
    """Use Opus to review all questions and select the best ones for a 2-page guide.

    Args:
        all_questions: Dict of {slide_num: [questions]}
        total_slides: Total number of slides

    Returns:
        Dict of {slide_num: [selected_questions]} with pre-selection flags
    """
    # Build a summary of all questions
    questions_summary = []
    for slide_num, qs in all_questions.items():
        for i, q in enumerate(qs):
            qtype = q.get("type", "unknown")
            if qtype == "open_ended":
                text = q.get("question", "")
            elif qtype == "short_answer":
                text = q.get("prompt", "")
            elif qtype == "fill_in_blank":
                text = q.get("sentence", "")
            elif qtype == "true_false":
                text = q.get("statement", "")
            elif qtype == "multiple_choice":
                text = q.get("question", "")
            elif qtype == "put_in_order":
                text = q.get("instruction", "")
            else:
                text = str(q)
            questions_summary.append({
                "slide": slide_num,
                "index": i,
                "type": qtype,
                "text": text[:200]  # Truncate for prompt
            })

    prompt = f"""You are reviewing questions for a student note-taking guide. The guide should be approximately 2 pages (maximum 3 pages).

There are {total_slides} slides total. Not every slide needs a question - select the BEST questions that:
1. Cover the most important concepts
2. Have good variety (mix of question types)
3. Are clear and well-written
4. Would help students engage with the material

IMPORTANT REQUIREMENTS:
1. 50-75% of selected questions should be "open_ended" or "short_answer" types (deeper engagement)
2. At least 25% should be OTHER types (fill_in_blank, true_false, multiple_choice, put_in_order) for variety

Question type priority for selection:
1. open_ended (highest priority - always good to include)
2. short_answer (high priority - tests recall of key terms/definitions)
3. fill_in_blank (medium priority)
4. true_false (medium priority)
5. multiple_choice (lower priority - only if conceptually meaningful)
6. put_in_order (use when appropriate for processes/sequences)

Here are all the generated questions:
{json.dumps(questions_summary, indent=2)}

Select approximately 15-20 questions total (enough for ~2 pages).
Remember: At least 50% must be open_ended or short_answer!

Return ONLY valid JSON listing which questions to INCLUDE:
{{
  "selected": [
    {{"slide": 4, "index": 0}},
    {{"slide": 4, "index": 2}},
    {{"slide": 5, "index": 1}},
    ...
  ]
}}"""

    response = client.messages.create(
        model=MODEL_ADVANCED,
        max_tokens=2000,
        messages=[{
            "role": "user",
            "content": prompt
        }]
    )

    response_text = response.content[0].text.strip()
    if response_text.startswith("```"):
        response_text = response_text.split("```")[1]
        if response_text.startswith("json"):
            response_text = response_text[4:]
        response_text = response_text.strip()

    try:
        result = json.loads(response_text)
        selected = result.get("selected", [])
        # Build a set of (slide, index) tuples for quick lookup
        selected_set = {(s["slide"], s["index"]) for s in selected}
        return selected_set
    except json.JSONDecodeError:
        # If parsing fails, select all questions
        return {(slide, i) for slide, qs in all_questions.items() for i in range(len(qs))}


def regenerate_questions(image, slide_num, selected_types, use_advanced_model=False):
    """Regenerate questions for a slide using only the selected question types."""
    base64_image = image_to_base64(image)
    model = MODEL_ADVANCED if use_advanced_model else MODEL_FAST

    type_instructions = []
    for qtype in selected_types:
        if qtype == "open_ended":
            type_instructions.append("- OPEN_ENDED: A thought-provoking question encouraging critical thinking")
        elif qtype == "short_answer":
            type_instructions.append("- SHORT_ANSWER: For acronyms, definitions, or terms (student writes brief answer)")
        elif qtype == "fill_in_blank":
            type_instructions.append("- FILL_IN_BLANK: A sentence with a key term blanked out with _____")
        elif qtype == "true_false":
            type_instructions.append("- TRUE_FALSE: A statement that is clearly true or false")
        elif qtype == "multiple_choice":
            type_instructions.append("- MULTIPLE_CHOICE: A question with 4 options (A-D), one correct")
        elif qtype == "put_in_order":
            type_instructions.append("- PUT_IN_ORDER: 3-5 items to arrange in correct sequence")

    types_str = "\n".join(type_instructions)
    num_questions = len(selected_types)

    prompt = f"""Analyze this lecture slide and generate exactly {num_questions} question(s) using ONLY these types:
{types_str}

Return ONLY valid JSON:
{{
  "questions": [
    // For open_ended (MUST include example_answer):
    {{"type": "open_ended", "question": "...", "example_answer": "A model answer (2-3 sentences)", "notes_prompt": "[Your notes:]"}},
    // For short_answer (for acronyms, definitions):
    {{"type": "short_answer", "prompt": "What does XYZ stand for?", "answer": "The full meaning"}},
    // For fill_in_blank:
    {{"type": "fill_in_blank", "sentence": "The _____ is...", "answer": "term"}},
    // For true_false:
    {{"type": "true_false", "statement": "...", "answer": true}},
    // For multiple_choice:
    {{"type": "multiple_choice", "question": "...", "options": ["A) ...", "B) ...", "C) ...", "D) ..."], "answer": "B"}},
    // For put_in_order:
    {{"type": "put_in_order", "instruction": "Arrange in order:", "items": ["...", "..."], "correct_order": [1, 0]}}
  ]
}}

IMPORTANT: For open_ended questions, always include an "example_answer" field with a model answer."""

    response = client.messages.create(
        model=model,
        max_tokens=1500,
        messages=[{
            "role": "user",
            "content": [
                {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": base64_image}},
                {"type": "text", "text": prompt}
            ]
        }]
    )

    response_text = response.content[0].text.strip()
    if response_text.startswith("```"):
        response_text = response_text.split("```")[1]
        if response_text.startswith("json"):
            response_text = response_text[4:]
        response_text = response_text.strip()

    try:
        return json.loads(response_text)
    except json.JSONDecodeError:
        return {"questions": [{"type": "open_ended", "question": response_text, "notes_prompt": "[Your notes:]"}]}


def toggle_slide_selection(slide_num, select_all):
    """Toggle all questions for a specific slide (callback for button on_click)."""
    qs = st.session_state.questions.get(slide_num, [])
    for i in range(len(qs)):
        st.session_state[f"slide_{slide_num}_q_{i}"] = select_all
    st.session_state.user_modified_selection = True


def on_checkbox_change():
    """Callback when user manually changes a checkbox."""
    st.session_state.user_modified_selection = True


def toggle_all_selection(select_all):
    """Toggle all questions across all slides (callback for global button)."""
    for sn in st.session_state.questions:
        for i in range(len(st.session_state.questions[sn])):
            st.session_state[f"slide_{sn}_q_{i}"] = select_all
    st.session_state.user_modified_selection = True

def format_question_display(q):
    """Format a structured question for display in the UI."""
    qtype = q.get("type", "open_ended")
    badge = f"[{QUESTION_TYPE_LABELS.get(qtype, qtype)}]"

    if qtype == "open_ended":
        return f"{badge} {q.get('question', '')}\n{q.get('notes_prompt', '[Your notes:]')}"

    elif qtype == "short_answer":
        return f"{badge} {q.get('prompt', '')}\n*Answer: {q.get('answer', '')}*"

    elif qtype == "fill_in_blank":
        return f"{badge} {q.get('sentence', '')}\n*Answer: {q.get('answer', '')}*"

    elif qtype == "true_false":
        answer = "True" if q.get("answer", False) else "False"
        return f"{badge} {q.get('statement', '')}\n○ True  ○ False\n*Answer: {answer}*"

    elif qtype == "multiple_choice":
        options = q.get("options", [])
        options_str = "  ".join(options)
        return f"{badge} {q.get('question', '')}\n{options_str}\n*Answer: {q.get('answer', '')}*"

    elif qtype == "put_in_order":
        items = q.get("items", [])
        items_str = " → ".join([f"[ {item} ]" for item in items])
        correct = q.get("correct_order", [])
        if correct and items:
            ordered = [items[i] for i in correct if i < len(items)]
            correct_str = " → ".join(ordered)
        else:
            correct_str = "N/A"
        return f"{badge} {q.get('instruction', 'Arrange in order:')}\n{items_str}\n*Correct order: {correct_str}*"

    return f"{badge} {str(q)}"


def set_cell_font(cell, font_name="Times New Roman", font_size=12):
    """Set font for a table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def format_question_for_docx(doc, q, question_num, show_answers=False):
    """Format a structured question for the Word document - clear visual hierarchy."""
    qtype = q.get("type", "open_ended")
    indent = Inches(0.3)  # Consistent indent for answer areas

    if qtype == "open_ended":
        # Question line
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(q.get('question', ''))
        p.paragraph_format.space_after = Pt(2)
        if show_answers:
            # Teacher version: show example answer
            example = q.get('example_answer', '')
            answer_p = doc.add_paragraph()
            answer_p.paragraph_format.left_indent = indent
            if example:
                run = answer_p.add_run("Example answer: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 100, 0)
                run2 = answer_p.add_run(example)
                run2.font.color.rgb = RGBColor(0, 100, 0)
            else:
                run = answer_p.add_run("[Open-ended response]")
                run.italic = True
                run.font.color.rgb = RGBColor(0, 100, 0)
            answer_p.paragraph_format.space_after = Pt(10)
        else:
            # Student version: lines for notes
            for i in range(2):
                line = doc.add_paragraph()
                run = line.add_run("_" * 70)
                run.font.size = Pt(14)
                line.paragraph_format.left_indent = indent
                line.paragraph_format.space_before = Pt(0)
                line.paragraph_format.space_after = Pt(0) if i < 1 else Pt(10)

    elif qtype == "short_answer":
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(q.get('prompt', ''))
        p.paragraph_format.space_after = Pt(2)
        if show_answers:
            # Teacher version: show answer in bold green
            ans_p = doc.add_paragraph()
            ans_p.paragraph_format.left_indent = indent
            run = ans_p.add_run(f"Answer: {q.get('answer', '')}")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0)
            ans_p.paragraph_format.space_after = Pt(10)
        else:
            # Student version: single line for answer (same as open-ended)
            line = doc.add_paragraph()
            run = line.add_run("_" * 70)
            run.font.size = Pt(14)
            line.paragraph_format.left_indent = indent
            line.paragraph_format.space_after = Pt(10)

    elif qtype == "fill_in_blank":
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        answer = q.get('answer', '')
        if show_answers:
            # Teacher version: show answer in bold green
            sentence = q.get('sentence', '').replace('_____', f"[{answer}]")
            p.add_run(sentence)
            p.paragraph_format.space_after = Pt(2)
            ans_p = doc.add_paragraph()
            ans_p.paragraph_format.left_indent = indent
            run = ans_p.add_run(f"Answer: {answer}")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0)
            ans_p.paragraph_format.space_after = Pt(10)
        else:
            sentence = q.get('sentence', '').replace('_____', '_' * 20)
            p.add_run(sentence)
            p.paragraph_format.space_after = Pt(10)

    elif qtype == "true_false":
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(q.get('statement', ''))
        p.paragraph_format.space_after = Pt(2)
        answer = q.get('answer', True)
        if show_answers:
            # Teacher version: highlight correct answer
            opts = doc.add_paragraph()
            opts.paragraph_format.left_indent = indent
            if answer:
                run_t = opts.add_run("● True")
                run_t.bold = True
                run_t.font.color.rgb = RGBColor(0, 100, 0)
                opts.add_run("                    ○ False")
            else:
                opts.add_run("○ True                    ")
                run_f = opts.add_run("● False")
                run_f.bold = True
                run_f.font.color.rgb = RGBColor(0, 100, 0)
            opts.paragraph_format.space_after = Pt(10)
        else:
            opts = doc.add_paragraph("○  True                    ○  False")
            opts.paragraph_format.left_indent = indent
            opts.paragraph_format.space_after = Pt(10)

    elif qtype == "multiple_choice":
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(q.get('question', ''))
        p.paragraph_format.space_after = Pt(2)

        options = q.get("options", [])
        correct = q.get("answer", "")

        if options:
            table = doc.add_table(rows=1, cols=len(options))
            table.autofit = True
            row = table.rows[0]
            for i, opt in enumerate(options):
                cell = row.cells[i]
                para = cell.paragraphs[0]
                # Check if this is the correct answer
                is_correct = opt.startswith(correct) if correct else False
                if show_answers and is_correct:
                    run = para.add_run(opt)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 100, 0)
                else:
                    para.add_run(opt)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
            # Indent the table
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblInd = OxmlElement('w:tblInd')
            tblInd.set(qn('w:w'), str(int(indent.twips)))
            tblInd.set(qn('w:type'), 'dxa')
            tblPr.append(tblInd)
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)

            if show_answers:
                ans_p = doc.add_paragraph()
                ans_p.paragraph_format.left_indent = indent
                run = ans_p.add_run(f"Answer: {correct}")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 100, 0)
                ans_p.paragraph_format.space_after = Pt(10)
            else:
                p_after = doc.add_paragraph()
                p_after.paragraph_format.space_before = Pt(0)
                p_after.paragraph_format.space_after = Pt(10)

    elif qtype == "put_in_order":
        instruction = q.get("instruction", "Arrange in order:")
        items = q.get("items", [])
        items_text = ", ".join(items)
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(f"{instruction} {items_text}")
        p.paragraph_format.space_after = Pt(2)

        if show_answers:
            correct_order = q.get("correct_order", [])
            if correct_order and items:
                ordered = [items[i] for i in correct_order if i < len(items)]
                correct_text = " → ".join(ordered)
            else:
                correct_text = ", ".join(items)
            ans_p = doc.add_paragraph()
            ans_p.paragraph_format.left_indent = indent
            run = ans_p.add_run(f"Correct order: {correct_text}")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0)
            ans_p.paragraph_format.space_after = Pt(10)
        else:
            answer = doc.add_paragraph("Order: " + "_" * 45)
            answer.paragraph_format.left_indent = indent
            answer.paragraph_format.space_after = Pt(10)

    else:
        p = doc.add_paragraph()
        p.add_run(f"Question {question_num}. ").bold = True
        p.add_run(str(q))
        p.paragraph_format.space_after = Pt(10)



def create_canvas_qti(questions_by_slide, quiz_title="Note-Taking Quiz"):
    """Generate a Canvas-compatible QTI ZIP file for quiz import."""
    import zipfile
    import xml.etree.ElementTree as ET
    from xml.dom import minidom
    import uuid
    
    def make_id():
        return f"g{uuid.uuid4().hex[:16]}"
    
    # Create the assessment XML
    assessment_id = make_id()
    
    # QTI 1.2 format for Canvas
    qti_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<questestinterop xmlns="http://www.imsglobal.org/xsd/ims_qtiasiv1p2"
                 xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
                 xsi:schemaLocation="http://www.imsglobal.org/xsd/ims_qtiasiv1p2 http://www.imsglobal.org/xsd/ims_qtiasiv1p2p1.xsd">
  <assessment ident="{assessment_id}" title="{quiz_title}">
    <qtimetadata>
      <qtimetadatafield>
        <fieldlabel>qmd_timelimit</fieldlabel>
        <fieldentry></fieldentry>
      </qtimetadatafield>
    </qtimetadata>
    <section ident="root_section">
'''
    
    question_num = 0
    for slide_num, questions in questions_by_slide.items():
        for q in questions:
            question_num += 1
            qtype = q.get("type", "open_ended")
            q_id = make_id()
            
            if qtype == "multiple_choice":
                question_text = q.get("question", "")
                options = q.get("options", [])
                correct = q.get("answer", "")
                
                # Build response labels
                responses_xml = ""
                correct_id = ""
                for i, opt in enumerate(options):
                    opt_id = make_id()
                    # Check if this is the correct answer (e.g., "A)" matches "A")
                    if opt.startswith(correct):
                        correct_id = opt_id
                    responses_xml += f'''
              <response_label ident="{opt_id}">
                <material><mattext texttype="text/html">{opt}</mattext></material>
              </response_label>'''
                
                qti_xml += f'''
      <item ident="{q_id}" title="Question {question_num}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>multiple_choice_question</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material><mattext texttype="text/html">{question_text}</mattext></material>
          <response_lid ident="response1" rcardinality="Single">
            <render_choice>{responses_xml}
            </render_choice>
          </response_lid>
        </presentation>
        <resprocessing>
          <outcomes><decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/></outcomes>
          <respcondition continue="No">
            <conditionvar><varequal respident="response1">{correct_id}</varequal></conditionvar>
            <setvar action="Set" varname="SCORE">100</setvar>
          </respcondition>
        </resprocessing>
      </item>
'''
            
            elif qtype == "true_false":
                statement = q.get("statement", "")
                answer = q.get("answer", True)
                true_id = make_id()
                false_id = make_id()
                correct_id = true_id if answer else false_id
                
                qti_xml += f'''
      <item ident="{q_id}" title="Question {question_num}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>true_false_question</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material><mattext texttype="text/html">{statement}</mattext></material>
          <response_lid ident="response1" rcardinality="Single">
            <render_choice>
              <response_label ident="{true_id}">
                <material><mattext texttype="text/html">True</mattext></material>
              </response_label>
              <response_label ident="{false_id}">
                <material><mattext texttype="text/html">False</mattext></material>
              </response_label>
            </render_choice>
          </response_lid>
        </presentation>
        <resprocessing>
          <outcomes><decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/></outcomes>
          <respcondition continue="No">
            <conditionvar><varequal respident="response1">{correct_id}</varequal></conditionvar>
            <setvar action="Set" varname="SCORE">100</setvar>
          </respcondition>
        </resprocessing>
      </item>
'''
            
            elif qtype == "short_answer":
                prompt_text = q.get("prompt", "")
                answer = q.get("answer", "")
                
                qti_xml += f'''
      <item ident="{q_id}" title="Question {question_num}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>short_answer_question</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material><mattext texttype="text/html">{prompt_text}</mattext></material>
          <response_str ident="response1" rcardinality="Single">
            <render_fib><response_label ident="answer1"/></render_fib>
          </response_str>
        </presentation>
        <resprocessing>
          <outcomes><decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/></outcomes>
          <respcondition continue="No">
            <conditionvar><varequal respident="response1">{answer}</varequal></conditionvar>
            <setvar action="Set" varname="SCORE">100</setvar>
          </respcondition>
        </resprocessing>
      </item>
'''
            
            elif qtype == "fill_in_blank":
                sentence = q.get("sentence", "").replace("_____", "[blank]")
                answer = q.get("answer", "")
                
                qti_xml += f'''
      <item ident="{q_id}" title="Question {question_num}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>short_answer_question</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material><mattext texttype="text/html">{sentence}</mattext></material>
          <response_str ident="response1" rcardinality="Single">
            <render_fib><response_label ident="answer1"/></render_fib>
          </response_str>
        </presentation>
        <resprocessing>
          <outcomes><decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/></outcomes>
          <respcondition continue="No">
            <conditionvar><varequal respident="response1">{answer}</varequal></conditionvar>
            <setvar action="Set" varname="SCORE">100</setvar>
          </respcondition>
        </resprocessing>
      </item>
'''
            
            elif qtype == "open_ended":
                question_text = q.get("question", "")
                
                qti_xml += f'''
      <item ident="{q_id}" title="Question {question_num}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>essay_question</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material><mattext texttype="text/html">{question_text}</mattext></material>
          <response_str ident="response1" rcardinality="Single">
            <render_fib><response_label ident="answer1" rshuffle="No"/></render_fib>
          </response_str>
        </presentation>
        <resprocessing>
          <outcomes><decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/></outcomes>
        </resprocessing>
      </item>
'''
            
            elif qtype == "put_in_order":
                instruction = q.get("instruction", "Arrange in order:")
                items = q.get("items", [])
                items_text = ", ".join(items)
                # Canvas doesn't have native ordering - convert to essay
                qti_xml += f'''
      <item ident="{q_id}" title="Question {question_num}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>essay_question</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material><mattext texttype="text/html">{instruction} {items_text}</mattext></material>
          <response_str ident="response1" rcardinality="Single">
            <render_fib><response_label ident="answer1" rshuffle="No"/></render_fib>
          </response_str>
        </presentation>
        <resprocessing>
          <outcomes><decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/></outcomes>
        </resprocessing>
      </item>
'''
    
    qti_xml += '''
    </section>
  </assessment>
</questestinterop>'''
    
    # Create manifest XML
    resource_id = make_id()
    manifest_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<manifest identifier="{make_id()}" xmlns="http://www.imsglobal.org/xsd/imsccv1p1/imscp_v1p1">
  <metadata>
    <schema>IMS Content</schema>
    <schemaversion>1.1.3</schemaversion>
  </metadata>
  <organizations/>
  <resources>
    <resource identifier="{resource_id}" type="imsqti_xmlv1p2" href="{assessment_id}.xml">
      <file href="{assessment_id}.xml"/>
    </resource>
  </resources>
</manifest>'''
    
    # Create ZIP file in memory
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('imsmanifest.xml', manifest_xml)
        zf.writestr(f'{assessment_id}.xml', qti_xml)
    
    buffer.seek(0)
    return buffer

def create_docx(title, intro_summary, questions_by_slide, outro_summary, show_answers=False):
    """Generate a compact Word document - Times New Roman, narrow margins, no title."""
    doc = Document()

    # Set narrow margins (0.5 inch)
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # Set default font to Times New Roman 10pt
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # No title - just start with content

    if intro_summary:
        p = doc.add_paragraph()
        run = p.add_run("Overview: ")
        run.bold = True
        p.add_run(intro_summary)
        p.paragraph_format.space_after = Pt(8)

    # Number questions sequentially across all slides
    question_num = 1
    for slide_num, questions in questions_by_slide.items():
        for q in questions:
            format_question_for_docx(doc, q, question_num, show_answers)
            question_num += 1

    if outro_summary:
        p = doc.add_paragraph()
        run = p.add_run("Key Takeaways: ")
        run.bold = True
        p.add_run(outro_summary)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- STREAMLIT UI ---

st.set_page_config(page_title="Slide Guide Generator", layout="wide")
st.title("📚 Note-Taking Guide Generator")
st.write("Upload a presentation PDF to generate a student note-taking guide.")

# Initialize session state
if "analyzed" not in st.session_state:
    st.session_state.analyzed = False
if "slides" not in st.session_state:
    st.session_state.slides = []
if "questions" not in st.session_state:
    st.session_state.questions = {}
if "include_intro" not in st.session_state:
    st.session_state.include_intro = True
if "include_outro" not in st.session_state:
    st.session_state.include_outro = True

# File upload
uploaded_file = st.file_uploader("Upload PDF", type="pdf")

if uploaded_file and not st.session_state.analyzed:
    # Read file bytes for hashing (need to do this before other reads)
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)  # Reset for later use
    file_hash = get_file_hash(file_bytes)

    # Check for cached results
    cached_data = load_from_cache(file_hash)

    if cached_data:
        st.info(f"📦 Found cached analysis for this file! (Hash: {file_hash[:8]}...)")
        use_cache_col1, use_cache_col2 = st.columns(2)
        with use_cache_col1:
            if st.button("✅ Use Cached Results"):
                with st.spinner("Loading cached results..."):
                    cleanup_working_dir()
                    images = convert_from_bytes(file_bytes, dpi=300)
                    for i, img in enumerate(images):
                        save_slide_image(img, i + 1)
                    st.session_state.slides = list(range(1, len(images) + 1))
                    st.session_state.slide_images = images
                    st.session_state.source_filename = uploaded_file.name
                    st.session_state.intro_summary = cached_data.get("intro_summary")
                    st.session_state.outro_summary = cached_data.get("outro_summary")
                    st.session_state.questions = {int(k): v for k, v in cached_data.get("questions", {}).items()}
                    # Restore selection state from cache
                    cached_selected = cached_data.get("selected", [])
                    selected_set = {(s, i) for s, i in cached_selected} if cached_selected else None
                    for slide_num, qs in st.session_state.questions.items():
                        for i in range(len(qs)):
                            if selected_set is not None:
                                st.session_state[f"slide_{slide_num}_q_{i}"] = (slide_num, i) in selected_set
                            else:
                                st.session_state[f"slide_{slide_num}_q_{i}"] = True
                st.session_state.analyzed = True
                st.session_state.ai_already_selected = True
                st.session_state.user_modified_selection = False
                st.rerun()
        with use_cache_col2:
            if st.button("🔄 Re-analyze (ignore cache)"):
                # Clear ALL caches including per-slide image cache
                clear_cache()
                st.session_state["force_reanalyze"] = True
                st.rerun()

    if not cached_data or st.session_state.get("force_reanalyze", False):
        # Pre-analysis options
        st.subheader("Analysis Options")
        col1, col2 = st.columns(2)
        with col1:
            generate_intro = st.checkbox("Generate introduction summary", value=True, key="gen_intro")
        with col2:
            generate_outro = st.checkbox("Generate conclusion summary", value=True, key="gen_outro")

        if st.button("🔍 Analyze Slides"):
            st.session_state.pop("force_reanalyze", None)  # Clear flag

            with st.spinner("Converting PDF to images..."):
                cleanup_working_dir()  # Start fresh
                images = convert_from_bytes(file_bytes, dpi=300)
                # Save full-size images to disk
                for i, img in enumerate(images):
                    save_slide_image(img, i + 1)
                st.session_state.slides = list(range(1, len(images) + 1))
                st.session_state.slide_images = images
                st.session_state.source_filename = uploaded_file.name
                total = len(images)

            st.write(f"Found {total} slides")

            # Analyze intro slides (first 2-3) if requested
            if generate_intro:
                with st.spinner("Analyzing introduction..."):
                    intro_texts = []
                    for i in range(min(3, total)):
                        intro_texts.append(analyze_slide(images[i], i+1, total, "intro"))
                    st.session_state.intro_summary = "\n".join(intro_texts)
            else:
                st.session_state.intro_summary = None

            # Analyze outro slides (last 2-3) if requested
            if generate_outro:
                with st.spinner("Analyzing conclusion..."):
                    outro_texts = []
                    for i in range(max(0, total-3), total):
                        outro_texts.append(analyze_slide(images[i], i+1, total, "outro"))
                    st.session_state.outro_summary = "\n".join(outro_texts)
            else:
                st.session_state.outro_summary = None

            # Analyze content slides
            content_start = min(3, total)
            content_end = max(content_start, total - 3)

            # Track cache hits for reporting
            cache_hits = 0
            api_calls = 0

            progress = st.progress(0)
            for idx, i in enumerate(range(content_start, content_end)):
                slide_num = i + 1
                image = images[i]

                # Check per-slide image cache first
                cached_qs, image_hash = get_cached_questions_for_image(image)

                if cached_qs:
                    # Found matching image in cache!
                    st.session_state.questions[slide_num] = cached_qs
                    for i in range(len(cached_qs)):
                        st.session_state[f"slide_{slide_num}_q_{i}"] = True
                    cache_hits += 1
                else:
                    # Need to analyze this slide
                    with st.spinner(f"Analyzing slide {slide_num}..."):
                        result = analyze_slide(image, slide_num, total, "content")
                        questions = result.get("questions", [])
                        # Verify answers immediately
                        verified = verify_answers(image, questions)
                        st.session_state.questions[slide_num] = verified
                        for i in range(len(verified)):
                            st.session_state[f"slide_{slide_num}_q_{i}"] = True
                        # Save to per-slide cache
                        save_questions_for_image(image_hash, verified)
                        api_calls += 1

                progress.progress((idx + 1) / (content_end - content_start))

            # Report cache usage
            if cache_hits > 0:
                st.toast(f"♻️ Reused {cache_hits} slides from cache, analyzed {api_calls} new slides")
            else:
                st.toast(f"📊 Analyzed {api_calls} slides")

            # Use Opus to review and pre-select the best questions for ~2 pages
            with st.spinner("Reviewing questions with advanced model..."):
                selected_set = review_and_select_questions(st.session_state.questions, total)
                # Update checkbox states based on Opus selection
                for slide_num, qs in st.session_state.questions.items():
                    for i in range(len(qs)):
                        is_selected = (slide_num, i) in selected_set
                        st.session_state[f"slide_{slide_num}_q_{i}"] = is_selected

            selected_count = len(selected_set)
            st.toast(f"🎯 Pre-selected {selected_count} questions for ~2 page guide")
            st.session_state.ai_already_selected = True
            st.session_state.user_modified_selection = False

            # Save questions to CSV for future retrieval
            save_questions_to_csv(
                st.session_state.questions,
                uploaded_file.name,
                selected_set
            )

            # Save full file to cache (for exact file match speedup)
            # Convert selected_set to list for JSON serialization
            selected_list = [[s, i] for s, i in selected_set]
            cache_data = {
                "intro_summary": st.session_state.intro_summary,
                "outro_summary": st.session_state.outro_summary,
                "questions": st.session_state.questions,
                "selected": selected_list
            }
            save_to_cache(file_hash, cache_data)

            st.session_state.analyzed = True
            st.rerun()

# Show results and selection UI
if st.session_state.analyzed:
    # Top controls
    top_col1, top_col2, top_col3 = st.columns([4, 2, 1])
    with top_col1:
        st.success("Analysis complete! Select the questions you want to include:")
    with top_col2:
        # Disable button if AI already selected and user hasn't made changes
        ai_already = st.session_state.get("ai_already_selected", False)
        user_modified = st.session_state.get("user_modified_selection", False)
        button_disabled = ai_already and not user_modified
        
        if button_disabled:
            st.button("🤖 Let AI select questions", key="ai_select_top", disabled=True,
                      help="AI has already pre-selected questions. Modify your selection first to re-run AI selection.")
        else:
            if st.button("🤖 Let AI select questions", key="ai_select_top",
                         help="Use AI to select the best questions for a ~2 page guide"):
                with st.spinner("AI is selecting the best questions..."):
                    total_slides = len(st.session_state.slide_images)
                    selected_set = review_and_select_questions(st.session_state.questions, total_slides)
                    # Update checkbox states based on AI selection
                    for slide_num, qs in st.session_state.questions.items():
                        for i in range(len(qs)):
                            is_selected = (slide_num, i) in selected_set
                            st.session_state[f"slide_{slide_num}_q_{i}"] = is_selected
                    st.toast(f"🎯 AI selected {len(selected_set)} questions for ~2 page guide")
                    st.session_state.ai_already_selected = True
                    st.session_state.user_modified_selection = False
                st.rerun()
    with top_col3:
        if st.button("🔄 Start Over", key="start_over_top"):
            cleanup_working_dir()
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()

    # Show intro summary (if generated)
    if st.session_state.intro_summary:
        with st.expander("📖 Introduction Summary", expanded=False):
            st.write(st.session_state.intro_summary)

    # Question selection
    st.subheader("Select Questions")

    # Calculate global selection state
    total_questions = sum(len(qs) for qs in st.session_state.questions.values())
    selected_count = sum(
        1 for sn, qs in st.session_state.questions.items()
        for i in range(len(qs))
        if st.session_state.get(f"slide_{sn}_q_{i}", True)
    )

    # Global toggle button - shows state and cycles through (uses callback to track user modification)
    global_col1, global_col2 = st.columns([1, 5])
    with global_col1:
        if selected_count == total_questions:
            # All selected - green check, click to deselect all
            st.button("✅ All Selected", key="global_toggle", type="primary",
                     on_click=toggle_all_selection, args=(False,))
        elif selected_count == 0:
            # None selected - X, click to select all
            st.button("✗ None Selected", key="global_toggle",
                     on_click=toggle_all_selection, args=(True,))
        else:
            # Some selected - partial, click to select all
            st.button(f"☐ {selected_count}/{total_questions}", key="global_toggle",
                     on_click=toggle_all_selection, args=(True,))

    for slide_num, qs in st.session_state.questions.items():
        with st.expander(f"Slide {slide_num}", expanded=True):
            col1, col2 = st.columns([1, 3])
            with col1:
                st.image(st.session_state.slide_images[slide_num-1], width=200)

            with col2:
                # Calculate per-slide selection state
                slide_selected = sum(1 for i in range(len(qs)) if st.session_state.get(f"slide_{slide_num}_q_{i}", True))
                slide_total = len(qs)

                # Per-slide toggle button (uses callback to avoid Streamlit button timing issues)
                slide_btn_col1, slide_btn_col2 = st.columns([1, 4])
                with slide_btn_col1:
                    if slide_selected == slide_total:
                        st.button("✅", key=f"toggle_{slide_num}", help="All selected - click to deselect",
                                  on_click=toggle_slide_selection, args=(slide_num, False))
                    elif slide_selected == 0:
                        st.button("✗", key=f"toggle_{slide_num}", help="None selected - click to select all",
                                  on_click=toggle_slide_selection, args=(slide_num, True))
                    else:
                        st.button(f"☐ {slide_selected}/{slide_total}", key=f"toggle_{slide_num}", help="Some selected - click to select all",
                                  on_click=toggle_slide_selection, args=(slide_num, True))

                # Display each question with full text and type badge
                for i, q in enumerate(qs):
                    key = f"slide_{slide_num}_q_{i}"
                    display_text = format_question_display(q)
                    # Initialize checkbox state if not exists
                    if key not in st.session_state:
                        st.session_state[key] = True
                    st.checkbox(display_text, key=key, on_change=on_checkbox_change)

            # View full image in a cleaner way - outside the columns, full width
            with st.popover("📷 View Full Size Slide"):
                image_path = get_slide_image_path(slide_num)
                if image_path.exists():
                    st.image(str(image_path), width="stretch")

            st.markdown("---")

            # Regeneration controls
            regen_key = f"regen_expand_{slide_num}"
            if st.button("🔄 Regenerate with...", key=f"regen_btn_{slide_num}"):
                st.session_state[regen_key] = not st.session_state.get(regen_key, False)

            if st.session_state.get(regen_key, False):
                st.write("Select question types to generate:")
                regen_types = []
                type_cols = st.columns(6)
                for idx, (qtype, label) in enumerate(QUESTION_TYPE_LABELS.items()):
                    with type_cols[idx]:
                        if st.checkbox(label, key=f"regen_type_{slide_num}_{qtype}"):
                            regen_types.append(qtype)

                # Model toggle
                use_opus = st.checkbox("🧠 Use advanced model (Opus)", key=f"use_opus_{slide_num}",
                                       help="Use Claude Opus for higher quality questions (slower, more expensive)")

                if st.button("Generate", key=f"do_regen_{slide_num}"):
                    if regen_types:
                        model_name = "Opus" if use_opus else "Sonnet"
                        with st.spinner(f"Regenerating questions with {model_name}..."):
                            image = st.session_state.slide_images[slide_num-1]
                            result = regenerate_questions(image, slide_num, regen_types, use_advanced_model=use_opus)
                            new_questions = result.get("questions", [])
                            # Verify answers for regenerated questions
                            verified = verify_answers(image, new_questions)
                            st.session_state.questions[slide_num] = verified
                            for i in range(len(verified)):
                                st.session_state[f"slide_{slide_num}_q_{i}"] = True
                            st.session_state[regen_key] = False
                            # Save regenerated questions to CSV
                            save_questions_to_csv(
                                {slide_num: verified},
                                st.session_state.get("source_filename", "unknown"),
                                None  # All newly regenerated questions considered selected
                            )
                        st.rerun()
                    else:
                        st.warning("Please select at least one question type.")

    # Show outro summary (if generated)
    if st.session_state.outro_summary:
        with st.expander("🎯 Conclusion Summary", expanded=False):
            st.write(st.session_state.outro_summary)
    
    # Download options
    st.subheader("Export Options")

    # Include toggles for summaries and LMS export
    export_col1, export_col2, export_col3 = st.columns(3)
    with export_col1:
        if st.session_state.intro_summary:
            st.session_state.include_intro = st.checkbox(
                "Include introduction summary",
                value=st.session_state.include_intro,
                key="export_include_intro"
            )
    with export_col2:
        if st.session_state.outro_summary:
            st.session_state.include_outro = st.checkbox(
                "Include conclusion summary",
                value=st.session_state.include_outro,
                key="export_include_outro"
            )
    with export_col3:
        export_canvas = st.checkbox(
            "Export for Canvas LMS",
            value=False,
            key="export_canvas",
            help="Generate a QTI file for importing into Canvas quizzes"
        )

    # Collect selected questions (read from checkbox keys)
    final_questions = {}
    for slide_num, qs in st.session_state.questions.items():
        selected_qs = [q for i, q in enumerate(qs) if st.session_state.get(f"slide_{slide_num}_q_{i}", True)]
        if selected_qs:
            final_questions[slide_num] = selected_qs

    # Determine what to include in export
    export_intro = st.session_state.intro_summary if st.session_state.include_intro else None
    export_outro = st.session_state.outro_summary if st.session_state.include_outro else None

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        student_docx = create_docx(
            "Lecture Notes",
            export_intro,
            final_questions,
            export_outro,
            show_answers=False
        )
        st.download_button(
            "📄 Student Version",
            student_docx,
            "note_guide_student.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with col2:
        # Check if any open-ended questions need example answers
        needs_examples = any(
            q.get("type") == "open_ended" and not q.get("example_answer")
            for qs in final_questions.values()
            for q in qs
        )
        
        if needs_examples:
            with st.spinner("Generating example answers for teacher guide..."):
                teacher_questions = generate_example_answers(
                    final_questions,
                    st.session_state.slide_images
                )
                # Save updated questions with example answers to CSV
                save_questions_to_csv(
                    teacher_questions,
                    st.session_state.get("source_filename", "unknown"),
                    None
                )
        else:
            teacher_questions = final_questions
        
        teacher_docx = create_docx(
            "Lecture Notes",
            export_intro,
            teacher_questions,
            export_outro,
            show_answers=True
        )
        st.download_button(
            "📝 Teacher Answer Key",
            teacher_docx,
            "note_guide_teacher.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with col3:
        if st.button("🗑️ Clean up images"):
            cleanup_working_dir()
            st.success("Working images cleaned up!")

    with col4:
        if st.button("🗄️ Clear analysis cache"):
            clear_cache()
            # Also reset session state so user re-analyzes with fresh data
            st.session_state.analyzed = False
            st.session_state.questions = {}
            st.success("Analysis cache cleared! Please re-analyze to use updated question generation.")
            st.rerun()

    # Canvas LMS export (only if checkbox is selected)
    if st.session_state.get("export_canvas", False):
        st.markdown("---")
        st.markdown("**Canvas LMS Export**")
        canvas_col1, canvas_col2 = st.columns([1, 3])
        with canvas_col1:
            canvas_qti = create_canvas_qti(final_questions, "Note-Taking Quiz")
            st.download_button(
                "📚 Download Canvas Quiz (.zip)",
                canvas_qti,
                "canvas_quiz.zip",
                "application/zip",
                help="Import this ZIP file into Canvas: Quizzes → ⋮ → Import"
            )
        with canvas_col2:
            st.caption(
                "To import: In Canvas, go to **Quizzes** → click **⋮** (kebab menu) → "
                "**Import** → select the ZIP file. Question types: Multiple Choice, "
                "True/False, Fill in Blank → Short Answer, Open-ended/Ordering → Essay."
            )

    # Questions history section
    if QUESTIONS_CSV_PATH.exists():
        st.markdown("---")
        st.markdown("**Questions History**")
        hist_col1, hist_col2 = st.columns([1, 3])
        with hist_col1:
            with open(QUESTIONS_CSV_PATH, 'rb') as f:
                st.download_button(
                    "📊 Download Questions History (CSV)",
                    f,
                    "questions_history.csv",
                    "text/csv",
                    help="All generated questions across all sessions"
                )
        with hist_col2:
            st.caption(
                "This CSV contains all questions generated across sessions, including question type, "
                "slide number, answers, and whether each question was selected. Import into Excel or "
                "Google Sheets for analysis."
            )

    # Reset button
    if st.button("🔄 Start Over"):
        cleanup_working_dir()
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()
